import io
from fastapi import UploadFile
from PIL import Image
from pydantic import BaseModel, validator
from docx import Document
from docx.shared import Inches


class UserRegister(BaseModel):
    login: str
    password: str

    @validator('password')
    def validate_password(cls, value):
        if len(value) < 8:
            raise ValueError("Пароль должен быть не менее 8 символов")
        return value

async def is_valid_image(file: UploadFile) -> bool:
    try:
        image = Image.open(io.BytesIO(await file.read()))
        image.verify()  # Проверяем, что файл является изображением
        return True
    except Exception:
        return False
    
async def create_docx_with_screenshots_and_comments(screenshots, comments, output_path):
    """
    Создает DOCX-файл, содержащий скриншоты и комментарии в таблице.

    :param screenshots: Список путей к файлам скриншотов.
    :param comments: Список комментариев (по одному для каждого скриншота).
    :param output_path: Путь для сохранения DOCX-файла.
    """
    # Создаем новый документ
    doc = Document()

    # Добавляем заголовок
    doc.add_heading('Скриншоты и комментарии', 0)

    # Создаем таблицу
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Добавляем заголовки столбцов
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Скриншот'
    hdr_cells[1].text = 'Комментарий'

    # Добавляем скриншоты и комментарии в таблицу
    for screenshot_path, comment in zip(screenshots, comments):
        row_cells = table.add_row().cells
        # Добавляем скриншот в первую ячейку
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(screenshot_path, width=Inches(2.0))  # Ширина скриншота 2 дюйма
        # Добавляем комментарий во вторую ячейку
        row_cells[1].text = comment

    # Сохраняем документ
    doc.save(output_path)